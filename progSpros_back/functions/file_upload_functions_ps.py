from copy import copy, deepcopy
from itertools import chain
import os
from typing import Any, Dict, Generator, List, Optional, Tuple, Union
from dataclasses import dataclass
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.parser import parse_xml
from docx.parts.hdrftr import FooterPart, HeaderPart
from docx.text.run import Run as DocxRun
from docxcompose.composer import Composer
from sqlalchemy import Column, Case
from openpyxl import load_workbook
from openpyxl.worksheet.worksheet import Worksheet
from openpyxl.cell import Cell
from openpyxl.styles import PatternFill
from openpyxl.utils.cell import coordinate_to_tuple
from openpyxl.utils import get_column_letter
import re
from flask import current_app
from io import BytesIO

@dataclass
class CaseDescriptor:
    sql_case: Case
    join: Optional[Tuple] = None

@dataclass
class ColumnDescriptor:
    db_column: Optional[Column] = None
    case_desc: Optional[CaseDescriptor] = None
    excel_title: Optional[str] = None
    mapping: Optional[Dict] = None
    aggr_func: Optional[callable] = None
    is_filter: Optional[bool] = False
    template_col_id: Optional[str] = None


@dataclass
class TableDescriptor:
    list_name: str # новое название
    data: List
    location: Optional[str] = 'A1' # координата верхнего левого угла
    main_cols: Optional[Dict[str, ColumnDescriptor]] = None
    optional_cols: Optional[Dict[str, ColumnDescriptor]] = None
    groupings_headers_height: int = 0
    highlight_pattern: Optional[str] = None
    highlight_color: Optional[str] = 'ffff00'
    highlighted_cols: Optional[List[str]] = None
    row_highlight_col: Optional[str] = None
    row_highlight_pattern: Optional[str] = None

@dataclass
class TableBlockDescriptor:
    template_ws : Worksheet
    col_num : Optional[int] = None
    value : Optional[str] = None

@dataclass(frozen=True)
class HashableDict:
    data: tuple

    @classmethod
    def from_dict(cls, d: dict):
        return cls(tuple(sorted(d.items())))
    
    def to_dict(self):
        return dict(self.data)
    

def prepare_all_data(titles: Dict[str, ColumnDescriptor], rows: List[Dict]):
    parted_rows = {}
    years_range = (100500, 0)
    
    
    titles['start_year'] = ColumnDescriptor(excel_title='Начало отбора')

    def get_row_key_dict(row: Dict) -> Dict:
        key_dict = HashableDict.from_dict({k: v for k, v in row.items() if k not in ['sum', 'year']})
        return key_dict
    
    for row in rows:
        
        year = row.get('year')
        if year:
            titles[year] = ColumnDescriptor(excel_title=str(year))
            years_range = (min(years_range[0], year), max(years_range[1], year))
        key = get_row_key_dict(row)
        if not parted_rows.get(key):
            
            new_row = {'sum': [row.get('sum')], 'start_year': [year]}
            new_row[year] = [row.get('sum')]
            parted_rows[key] = new_row
        else:
            new_row = parted_rows[key]
            new_row['sum'].append(row.get('sum'))
            new_row['start_year'].append(year)
            if year not in new_row:
                new_row[year] = []
            new_row[year].append(row.get('sum'))
            parted_rows[key] = new_row
    result = []
    for start, end in parted_rows.items():
        end['sum'] = sum(filter((lambda x : x is not None), end['sum']))
        end['start_year'] = min(filter((lambda x : x is not None), end['start_year']))
        if None in end:
            end.pop(None)
        result_row = copy(start.to_dict())
        result_row['sum'] = float(end['sum'])
        result_row['start_year'] = end['start_year']
        end.pop('sum')
        end.pop('start_year')

        for year, value in end.items():
            result_row[year] = float(sum(filter((lambda x : x is not None), value)))
        
        result.append(result_row)
    
    return (result, years_range)

class UtilsMixin:
    @classmethod
    def _get_template_path(cls, template_name: str) -> str:
        """
        Шаблон лежит в app/templates.
        current_app.root_path указывает на папку app.
        """
        return os.path.join(current_app.root_path, "templates", template_name)

class ExcelBuilder(UtilsMixin):

    def _build_templates_mapper(self):
        self.templates_mapper = {}
        if self.templates_ws is None:
            return
        
        cur_cell : Cell = self.templates_ws[f'A{self.table_desc.groupings_headers_height+1}']

        while cur_cell.value is not None:
            self.templates_mapper[cur_cell.value] = self.templates_ws.cell(column=cur_cell.column, row=1).column
            cur_cell = cur_cell.offset(column=1)
    
    
    def _build_optionals_mapper(self):
        self.optional_mapper = {}
        if self.optional_ws is None:
            return
        
        cur_cell : Cell = self.optional_ws[f'A{self.table_desc.groupings_headers_height+1}']
        
        while cur_cell.value is not None:
            self.optional_mapper[cur_cell.value] = self.optional_ws.cell(column=cur_cell.column, row=1).column
            cur_cell = cur_cell.offset(column=1)

    def _build_headers_tree(self, template_cols: List[TableBlockDescriptor]) -> Dict:
        tree_height = self.table_desc.groupings_headers_height + 1
        tree = {'_start': {}, '_col_width': 0, '_prev': None}
        # строим дерево
        for cur_col_desc in template_cols:
            cell = cur_col_desc.template_ws[f'{get_column_letter(cur_col_desc.col_num)}1']
            cur_lvl = tree['_start']
            
            for i in range(tree_height - 1):
                cur_title = str(cell.value)
                if cur_title not in cur_lvl:
                    cur_lvl[cur_title] = {'_col_width': 0, '_prev': cur_lvl}
                cur_lvl = cur_lvl[cur_title]
                
                cell = cell.offset(row=1)
            
            cur_lvl[str(cell.value) if cur_col_desc.value is None else str(cur_col_desc.value)] = cur_col_desc
            cur_lvl['_col_width'] = cur_lvl.get('_col_width', 1)

            while cur_lvl.get('_prev'):
                cur_lvl['_col_width'] += 1
                cur_lvl = cur_lvl['_prev']


        
        return tree

    def _get_headers_tree_iterator(self, tree: Dict, with_widths: Optional[bool] = False):
        def _traverse(node):
            for key, value in node.items():

                if key.startswith('_'):
                    continue
                    
                if isinstance(value, TableBlockDescriptor):
                    yield value    
                elif isinstance(value, dict):
                    if with_widths:
                        yield (key, value.get('_col_width'))
                    yield from _traverse(value)
        
        return _traverse(tree.get('_start', {}))
    
    @staticmethod
    def _copy_cell_style(target_cell : Cell, source_cell : Cell):
        t_rows_height = getattr(source_cell.parent.row_dimensions[source_cell.row], 'height', None)

        # if t_rows_height is not None:
        target_cell.parent.row_dimensions[target_cell.row].height = t_rows_height
        target_cell._style = source_cell._style
        target_cell.number_format = source_cell.number_format
    

    def _copy_headers(self, headers_tree: Dict):
        header_iterator = self._get_headers_tree_iterator(headers_tree, with_widths=True)
        insert_col : int = self.main_ws[self.table_desc.location].column #!
        self.mappers_mapping = {}
        cells_to_merge = {}
        for header_column in header_iterator:
            insert_row = self.main_ws[self.table_desc.location].row #!

            if isinstance(header_column, tuple):
                if (width := header_column[1]) != 1:
                    cells_to_merge[(insert_row, insert_col)] = width
                continue
            if header_column.template_ws is self.main_ws:
                if (cell_width := cells_to_merge.get((insert_row, insert_col), False)):
                    self.main_ws.merge_cells(start_row=insert_row, start_column=insert_col, end_row=insert_row, end_column=insert_col+cell_width-1)
                insert_col += 1
                continue
            source_ws = header_column.template_ws
            
            for row in range(1, self.table_desc.groupings_headers_height + 2):
                source_cell = source_ws.cell(row, header_column.col_num)
                
                target_cell = self.main_ws.cell(column=insert_col, row=insert_row)
                
                try:
                    self._copy_cell_style(target_cell=target_cell, source_cell=source_cell)
                    target_cell.parent.row_dimensions[target_cell.row].height = None
                    target_cell.value = source_cell.value
                except AttributeError:
                    pass
                if (cell_width := cells_to_merge.get((insert_row, insert_col), False)):
                    self.main_ws.merge_cells(start_row=insert_row, start_column=insert_col, end_row=insert_row, end_column=insert_col+cell_width-1)
                insert_row += 1
            if header_column.value:
                    target_cell.value = header_column.value
            else:
                target_cell.value = source_cell.value

            row += 1
            source_cell = source_ws.cell(row, header_column.col_num)
            target_cell = self.main_ws.cell(column=insert_col, row=insert_row)
            self._copy_cell_style(target_cell=target_cell, source_cell=source_cell)
            target_cell.parent.column_dimensions[get_column_letter(target_cell.column)] = target_cell.parent.column_dimensions[get_column_letter(source_cell.column)]

            insert_col += 1
            
    def _build_main_template(self, data_cols : Optional[List[str]] = None):
        existing_cols_ids : List['str'] = list(self.table_desc.main_cols.keys())
        if self.table_desc.optional_cols is not None and data_cols is not None:
            existing_cols_ids.extend(set(self.table_desc.optional_cols.keys()).intersection(set(data_cols)))
        main_cols_titles = set()
        template_cols : List[TableBlockDescriptor] = []

        cur_cell : Cell = self.main_ws[self.table_desc.location]
        cur_cell = cur_cell.offset(row=self.table_desc.groupings_headers_height)
        while cur_cell.value is not None:
            main_cols_titles.add(cur_cell.value)
            
            template_cols.append(TableBlockDescriptor(template_ws=self.main_ws, col_num=cur_cell.col_idx))
            cur_cell = cur_cell.offset(column=1)

        
        
        for col_id in existing_cols_ids:
            col_desc = self.table_desc.main_cols.get(col_id)
            if col_desc is None and self.table_desc.optional_cols and self.optional_ws:
                col_desc = self.table_desc.optional_cols.get(col_id)
            if not getattr(col_desc, 'excel_title', False):
                continue
            
            if col_desc.excel_title in main_cols_titles:
                continue

            if col_desc.template_col_id:
                col_num = self.templates_mapper.get(col_desc.template_col_id)
                value = col_desc.excel_title
                template_cols.append(TableBlockDescriptor(template_ws=self.templates_ws, col_num=col_num, value=value))
                
                continue
            if (col_num := self.optional_mapper.get(col_desc.excel_title)) is not None:
                template_cols.append(TableBlockDescriptor(template_ws=self.optional_ws, col_num=col_num))

        headers_tree = self._build_headers_tree(template_cols)
        self._copy_headers(headers_tree)

    def _write_data(self):
        highlight_fill = None
        if (self.highlight_pattern is not None or self.row_highlight_pattern is not None):
            highlight_fill = PatternFill(start_color=self.highlight_color, end_color=self.highlight_color, fgColor=self.highlight_color, fill_type="solid")
        if (self.highlight_pattern is not None):
            self.highlight_pattern = re.compile(self.highlight_pattern)
            cols_to_highlight = set()
        if (self.row_highlight_pattern is not None):
            self.row_highlight_pattern = re.compile(self.row_highlight_pattern)
        
        max_col = self.main_ws.max_column
        self.start_row, self.start_column = coordinate_to_tuple(self.table_desc.location) 
        self.start_row += self.table_desc.groupings_headers_height
        
        self.cols_mapper = {
            col.excel_title: key 
            for key, col in chain(
                self.table_desc.main_cols.items(),
                self.table_desc.optional_cols.items() if self.table_desc.optional_cols else []
            )
        }
        template_headers_mapper = {}
        
        for col in range(self.start_column, max_col + 1):
            if self.main_ws.cell(self.start_row, col).value is None:
                continue
            col_id = self.cols_mapper.get(str(self.main_ws.cell(self.start_row, col).value).replace('\n', ' '))
            
            template_headers_mapper[col_id] = col 
            
            if col_id in self.highlighted_cols:
                cols_to_highlight.add(col)
        self.start_row += 1
        template_cells = [self.main_ws.cell(self.start_row, c) for c in range(1, max_col + 1)]
        template_styles = [c._style for c in template_cells]
        template_numfmts = [c.number_format for c in template_cells]
        
        rows_touched = set()
        del_cells = []
        for (r, col) in self.main_ws._cells.keys():
            if r > self.start_row and 1 <= col <= max_col:
                del_cells.append(self.main_ws._cells[(r, col)])
                rows_touched.add(r)

        for cell in del_cells:
            del cell

        for r in rows_touched:
            if r in self.main_ws.row_dimensions:
                del self.main_ws.row_dimensions[r]

        for col in range(self.start_column, max_col + 1):
            self.main_ws.cell(self.start_row, col).value = None

        if not self.data or len(self.data) == 0:
            bio = BytesIO()
            self._wb.save(bio)
            return bio.getvalue()

        def apply_template_style(row_idx: int) -> None:
            self.main_ws.row_dimensions[row_idx].height = None
            for col_idx in range(self.start_column, max_col + 1):
                cell = self.main_ws.cell(row_idx, col_idx)
                cell._style = copy(template_styles[col_idx - 1])
                cell.number_format = template_numfmts[col_idx - 1]

        def write_row_values(row_idx: int, row: Dict[str, Any]) -> None:
            for cell_id, cell_value in row.items():
                col_idx = template_headers_mapper.get(cell_id)
                
                if not col_idx:
                    continue
                try:
                    self.main_ws.cell(row=row_idx, column=col_idx).value = cell_value
                    if self.highlight_pattern is not None and (self.highlighted_cols is None or col_idx in cols_to_highlight) and self.highlight_pattern.match(cell_value):
                        self.main_ws.cell(row=row_idx, column=col_idx).fill=highlight_fill  
                except Exception as e:
                    raise e

        for i, row in enumerate(self.data):
            row_idx = self.start_row + i
            apply_template_style(row_idx)
            write_row_values(row_idx, row)
            if self.row_highlight_pattern is not None and self.row_highlight_pattern.match(str(row.get(self.row_highlight_col))):
                for col_idx in range(self.start_column, max_col + 1):
                    self.main_ws.cell(row_idx, col_idx).fill = highlight_fill
                if self._clean_row_hl_checks:
                    hl_col_idx = template_headers_mapper.get(self.row_highlight_col)
                    if hl_col_idx:
                        cell = self.main_ws.cell(row_idx, hl_col_idx)
                        cell.value = self.row_highlight_pattern.sub('', str(cell.value))

    def build_export_xlsx(self, template_name: str, *table_descs, clean_row_hl_checks: bool = True):
        template_path = self._get_template_path(template_name)
        if not os.path.exists(template_path):
            raise ValueError(f"Не найден шаблон выгрузки: {template_path}")
        
        self._wb = load_workbook(template_path)
        self._clean_row_hl_checks = clean_row_hl_checks
        for table_desc in table_descs:
            self.table_desc = table_desc
            self.highlight_pattern = self.table_desc.highlight_pattern
            self.highlight_color = self.table_desc.highlight_color
            self.highlighted_cols = self.table_desc.highlighted_cols or []
            self.row_highlight_col = self.table_desc.row_highlight_col
            self.row_highlight_pattern = self.table_desc.row_highlight_pattern
            self.data = table_desc.data

            self.main_ws = self._wb[f'{table_desc.list_name}_main']
            self.optional_ws = self._wb[f'{table_desc.list_name}_optional'] if f'{table_desc.list_name}_optional' in self._wb.sheetnames else None
            self.templates_ws = self._wb[f'{table_desc.list_name}_template_cols'] if f'{table_desc.list_name}_template_cols' in self._wb else None

            self._build_templates_mapper()
            self._build_optionals_mapper()
            
            self._build_main_template(self.data[0].keys() if len(self.data) > 0 else None)

            self._write_data()

            self.main_ws.title = table_desc.list_name

            for ws in [self.templates_ws, self.optional_ws]:
                if ws is not None:
                    self._wb.remove(ws)
            table_coords = f"{get_column_letter(self.start_column)}{self.start_row-1}:{get_column_letter(self.start_column+len(self.cols_mapper)-2)}{self.start_row+len(self.data)}"
            filters = self.main_ws.auto_filter
            filters.ref = table_coords
            self.main_ws.auto_filter.add_sort_condition(table_coords)

        bio = BytesIO()
        self._wb.save(bio)
        return bio.getvalue()


_HL_START = ''
_HL_END = ''
_HL_PATTERN = re.compile(f'{re.escape(_HL_START)}(.*?){re.escape(_HL_END)}', re.DOTALL)


class _KeyReplacer:
    def __init__(self, p, key, value) -> None:
        self.p = p
        self.key = key
        self.value = value
        self.run_text = ""
        self.runs_indexes: List = []
        self.run_char_indexes: List = []
        self.runs_to_change: Dict = {}

    def _build_indexes_table(self) -> None:
        run_index = 0
        for run in self.p.runs:
            text_len = len(run.text)
            self.run_text += run.text
            self.runs_indexes += [run_index] * text_len
            self.run_char_indexes += [char_index for char_index in range(text_len)]
            run_index += 1

    def replace(self) -> None:
        self._build_indexes_table()
        index_to_replace = self.run_text.find(self.key)

        for i in range(len(self.key)):
            index = index_to_replace + i
            run_index = self.runs_indexes[index]
            run = self.p.runs[run_index]
            run_char_index = self.run_char_indexes[index]

            if not self.runs_to_change.get(run_index):
                self.runs_to_change[run_index] = [char for char in run.text]

            run_to_change: Dict = self.runs_to_change.get(run_index)
            if index == index_to_replace:
                run_to_change[run_char_index] = self.value
            else:
                run_to_change[run_char_index] = ""

        for index, text in self.runs_to_change.items():
            run = self.p.runs[index]
            run.text = "".join(text)


class DocxBuilder(UtilsMixin):
    # PLACEHOLDER_PATTERN = r'%([^%]+?)%'
    PLACEHOLDER = lambda self, key: f'%{key}%'

    @staticmethod
    def _rm_nonbreaks_hyphens(s: str) -> Optional[str]:
        if s is None: return None
        return s.replace('\u00a0', ' ').replace('\u00ad', '')

    def fill_docx_template(self, template_name: str, **kwargs) -> Document:
        self.doc = Document(self._get_template_path(template_name))

        for key, value in kwargs.items():
            placeholder = self.PLACEHOLDER(key)
            for p in self._get_all_paragraphs():
                self._replace_key(p, placeholder, self._rm_nonbreaks_hyphens(str(value)))

        # for p in self._get_all_paragraphs():
        #     self._expand_highlight_markers(p)

        self._strip_special_chars(self.doc)

        return self.doc

    @staticmethod
    def _strip_special_chars(doc: Document) -> None:
        NBSP = ' '
        SOFT_HYPHEN = '­'

        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        for tag in ('w:softHyphen', 'w:noBreakHyphen'):
            for elem in doc.element.body.findall(f'.//{tag}', nsmap):
                elem.getparent().remove(elem)

        for section in [doc] + list(doc.sections):
            # Обращение к .header/.footer при отсутствующем колонтитуле заставляет
            # python-docx создать пустой колонтитул (_Header/_Footer._get_or_add_definition
            # в docx/section.py) — пропускаем такие случаи, иначе у шаблонов без
            # колонтитулов появляется паразитный пустой колонтитул.
            sources = [section] if section is doc else [
                hf for hf in (section.header, section.footer) if hf._has_definition
            ]
            for src in sources:
                for p in src.paragraphs:
                    for run in p.runs:
                        run.text = run.text.replace(NBSP, ' ').replace(SOFT_HYPHEN, '')
                for table in src.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.text = run.text.replace(NBSP, ' ').replace(SOFT_HYPHEN, '')
        

    def _get_all_paragraphs(self) -> List[Any]:
        paragraphs = []
        paragraphs.extend(self._get_paragraphs_from(self.doc))

        for section in self.doc.sections:
            # .header/.footer auto-vivify an empty definition when none exists
            # (see the note in _strip_special_chars) - skip templates without one.
            if section.header._has_definition:
                paragraphs.extend(self._get_paragraphs_from(section.header))
            if section.footer._has_definition:
                paragraphs.extend(self._get_paragraphs_from(section.footer))

        return paragraphs
    
    @staticmethod
    def _get_paragraphs_from(item: Union[Any]) -> Generator:
        yield from item.paragraphs

        for table in item.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        yield paragraph
    
    @classmethod
    def _replace_key(cls, paragraph, key: str, value: str):
        if key not in paragraph.text:
            return

        replace_success = False
        for run in paragraph.runs:
            if key in run.text:
                run.text = run.text.replace(key, value)
                replace_success = True

        if not replace_success:
            cls._replace_broken_key(paragraph, key, value)

    @staticmethod
    def _expand_highlight_markers(paragraph) -> None:
        for run in list(paragraph.runs):
            if _HL_START not in run.text:
                continue
            parts = _HL_PATTERN.split(run.text)
            if len(parts) <= 1:
                continue
            original_xml = deepcopy(run._element)
            run.text = parts[0]
            prev_elem = run._element
            for i, part in enumerate(parts[1:]):
                new_xml = deepcopy(original_xml)
                new_run = DocxRun(new_xml, run._parent)
                new_run.text = part
                # split() with a capture group alternates: [plain, captured, plain, captured, ...]
                # so parts[1:] has captured (highlighted) at even indices (0, 2, 4...)
                new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW if i % 2 == 0 else None
                prev_elem.addnext(new_xml)
                prev_elem = new_xml

    @classmethod
    def _replace_broken_key(cls, paragraph, key: str, value: str):
        while key in paragraph.text:
            replacer = _KeyReplacer(paragraph, key=key, value=value)
            replacer.replace()

    @staticmethod
    def join_docs(*docs: Document) -> Document:
        base_doc : Document = docs[0]
        composer = Composer(base_doc)
        for doc in docs[1:]:
            composer.append(doc)

        return base_doc

    @staticmethod
    def copy_headers_footers(dst_doc: Document, src_doc: Document) -> None:
        # docxcompose переносит колонтитулы только при склейке многораздельных документов (Composer.fix_header_and_footers).
        dst_section = dst_doc.sections[0]
        src_section = src_doc.sections[0]
        composer = Composer(dst_doc)

        for attr, is_header in (
            ('header', True), ('first_page_header', True), ('even_page_header', True),
            ('footer', False), ('first_page_footer', False), ('even_page_footer', False),
        ):
            src_hf = getattr(src_section, attr)
            if not src_hf._has_definition:
                continue

            dst_hf = getattr(dst_section, attr)
            if dst_hf._has_definition:
                dst_hf._drop_definition()

            # add_relationship() (used by docxcompose itself) would wrap the copied
            # part as a generic opc.Part, which lacks the .element StoryPart needs -
            # so build a proper Header/FooterPart and copy its own relationships
            # (e.g. a letterhead image) via add_referenced_parts instead.
            src_part = src_hf._definition
            part_cls = HeaderPart if is_header else FooterPart
            content_type = CT.WML_HEADER if is_header else CT.WML_FOOTER
            partname_pattern = "/word/header%d.xml" if is_header else "/word/footer%d.xml"

            new_partname = dst_doc.part.package.next_partname(partname_pattern)
            new_element = parse_xml(src_part.blob)
            new_part = part_cls(new_partname, content_type, new_element, dst_doc.part.package)
            composer.add_referenced_parts(src_part, new_part, new_element)

            reltype = RT.HEADER if is_header else RT.FOOTER
            rid = dst_doc.part.relate_to(new_part, reltype)

            kind = src_hf._hdrftr_index
            if is_header:
                dst_section._sectPr.add_headerReference(kind, rid)
            else:
                dst_section._sectPr.add_footerReference(kind, rid)

    @staticmethod
    def save_file(doc: Document) -> BytesIO:
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream